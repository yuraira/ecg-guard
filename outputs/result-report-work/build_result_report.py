from __future__ import annotations

import hashlib
import html
import shutil
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from PIL import Image as PILImage
from PIL import ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
WORK = Path(__file__).resolve().parent
REFERENCE = WORK / "official-template.docx"
DELIVERABLES = ROOT / "deliverables"
DOCX_OUT = DELIVERABLES / "2026 오픈소스 개발자대회 결과보고서_접수번호미정(ECG Guard).docx"
PDF_OUT = DELIVERABLES / "2026 오픈소스 개발자대회 결과보고서_접수번호미정(ECG Guard).pdf"
ARCH_IMAGE = WORK / "architecture.png"

MALGUN = Path(r"C:\Windows\Fonts\malgun.ttf")
MALGUN_BOLD = Path(r"C:\Windows\Fonts\malgunbd.ttf")
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(166, 166, 166)
BLUE = "A5C9EB"
PALE_BLUE = "DAE9F7"

TEAM_NAME = "ECG Guard"
PROJECT_NAME = "ECG Guard 신뢰도 기반 심전도 인공지능 평가 플랫폼"
REPOSITORY = "https://github.com/yuraira/ecg-guard"
PUBLIC_DEMO = "https://ecg-guard.onrender.com"
RELEASE_URL = "https://github.com/yuraira/ecg-guard/releases/tag/baseline-v1"
WEIGHT_URL = "https://github.com/yuraira/ecg-guard/releases/download/baseline-v1/best_model.pt"
WEIGHT_SHA256 = "44a8ecc96f1ac084db2ef6921bf8e438c1130da6be140d7fc3ac7fe3ecfa2ead"
CONTAINER_SBOM_SHA256 = "43cc83f8399271bcc34d82975fecdf7e199f95d7e946fe451764b722abd3abf4"

INTRO = (
    "공개 12유도 심전도에서 5개 상위 진단군의 다중 라벨 확률을 산출하고, "
    "신호 품질 경고와 예측 불확실성을 분리해 제시하는 연구·교육용 오픈소스 평가 플랫폼이다. "
    "환자 단위 분할, 검증 세트 임계값 고정, 하위집단·오류 분석, 공개 가중치와 SBOM으로 재현성과 투명성을 강화했다."
)

BACKGROUND = (
    "심전도 인공지능이 예측값만 보여주면 잡음이 큰 입력이나 결정 경계에 가까운 사례에서도 결과가 과신될 수 있다. "
    "또한 전체 평균 성능만으로는 진단군과 인구집단별 약점을 확인하기 어렵다. ECG Guard는 공개 PTB-XL을 이용해 "
    "(1) 재현 가능한 기준 모델을 만들고, (2) 입력 신호 품질과 모델 불확실성을 별도 검토 사유로 표시하며, "
    "(3) 성능·보정·오류·하위집단 결과와 모델 한계를 함께 공개하는 것을 목표로 했다. 본 결과는 의료 진단이나 의료기기 성능의 근거가 아니다."
)

ENVIRONMENT = (
    "하드웨어: NVIDIA GeForce RTX 4070 SUPER 학습, CPU 전용 배포 검증. "
    "소프트웨어: Windows/Linux, Python 3.12, PyTorch 2.12.1, NumPy 2.5.1, pandas 3.0.5, "
    "scikit-learn 1.9.0, WFDB 4.3.1, Streamlit 1.60.0. "
    "도구: Git/GitHub Actions, pytest 9.1.1, Docker Desktop/Engine 29.6.2, CycloneDX와 Syft 1.49.0."
)

ARCHITECTURE = (
    "PTB-XL 100 Hz WFDB 파형과 메타데이터를 읽어 공식 strat_fold 1-8/9/10을 각각 학습/검증/시험으로 고정한다. "
    "학습 fold에서만 산출한 정규화 통계를 적용한 뒤 residual 1D CNN이 NORM, MI, STTC, CD, HYP 확률을 출력한다. "
    "검증 fold에서 temperature와 진단군별 임계값, 선택적 판정 cutoff를 고정하고 시험 fold에서 한 번 평가한다. "
    "신호 품질 점수와 결정 경계 기반 uncertainty는 서로 다른 review 사유로 합성되며, CLI/Streamlit/JSON 출력이 동일한 추론 자원을 사용한다."
)

FEATURES = (
    "핵심 기능: 12유도 10초 ECG 입력 검증, 5개 진단군 다중 라벨 추론, temperature scaling 확률 보정, "
    "평탄 신호·극단 진폭·기준선 변동 등 설명 가능한 SQI 경고, 결정 경계 기반 불확실성 검토 라우팅, "
    "연령·성별·artifact 주석별 성능 분석, 결과 JSON 다운로드와 세션 삭제, 합성 ECG 시연을 제공한다.\n"
    "개발 과정: 데이터 무결성 및 환자 중복 0명 확인 -> residual 1D CNN 학습 -> 검증 fold에서 보정·임계값 잠금 -> "
    "시험 fold 평가와 patient-cluster bootstrap -> 신호 품질·선택적 판정·오류 분석 -> CLI/웹/Docker/Release/SBOM 공개 순으로 진행했다.\n"
    "구동·검증: pip 설치 후 CLI 또는 Streamlit을 실행할 수 있고 Docker는 비특권 사용자·읽기 전용 루트·512 MB 제한에서 확인했다. "
    "가중치는 baseline-v1 Release에서 인증 없이 내려받아 23,579,661 bytes와 SHA-256을 이중 검증한다. 자동화 테스트 66개, "
    "GitHub CI, 패키지 컴파일, pip check, Docker healthcheck가 통과했다. 공개 데모(https://ecg-guard.onrender.com)의 "
    "HTTPS healthcheck와 합성 분석·세션 삭제·잘못된 입력 오류 처리도 확인했다.\n"
    "시험 결과(2,158건, 환자 1,877명): macro AUROC 0.91627 (95% patient-cluster bootstrap 0.90871-0.92378), "
    "macro AUPRC 0.81057 (0.79269-0.82625), 민감도 0.82693, 특이도 0.85176, Brier 0.08652. "
    "80% 목표 선택적 판정의 실제 coverage는 77.43%이며 Hamming 오류율은 전체 0.1504에서 0.1270으로 감소했다. "
    "이는 선택된 하위집단의 trade-off이지 모델 자체 성능 향상이나 임상 안전성의 증거가 아니다."
)

IMPACT = (
    "분류 확률과 함께 입력 기술 상태·불확실성·검토 필요 사유를 제시해 의료 AI 결과를 비판적으로 해석하는 교육과 연구에 활용할 수 있다. "
    "고정된 평가 프로토콜, 공개 가중치, 모델 카드, CycloneDX SBOM, Docker 환경을 제공하므로 재현 실험과 후속 비교의 기반이 된다. "
    "향후 다른 공개 ECG 데이터셋 외부 검증, ensemble/OOD 탐지, 임상 전문가 사용성 평가, 세부 리듬 분류, 경량 추론으로 확장할 수 있다."
)

OTHER = (
    "차별성: 단순 분류가 아니라 보정·SQI·uncertainty·판정 보류·하위집단·오류 갤러리를 하나의 검증 파이프라인과 웹 화면에 연결했다. "
    "시험 fold를 개발 중 반복 사용하지 않고 검증 fold에서 의사결정을 잠근 점, 약점과 사용 금지 범위를 모델 카드에 공개한 점이 핵심이다.\n"
    "한계: 단일 데이터셋·단일 seed이며 외부 병원/장비/인종 검증이 없다. HYP precision은 0.39439로 false positive가 많다. "
    "SQI의 시험 artifact 오류 탐지 AUROC는 0.5568로 낮아 자동 배제 도구로 사용할 수 없다. uncertainty는 epistemic uncertainty가 아니라 결정 경계 근접도다.\n"
    "로드맵: 외부 데이터 잠금 평가 -> 반복 seed/ensemble -> OOD·품질 모델 고도화 -> 임상 전문가와 화면 표현 평가 -> 의료기기·FHIR 연동은 별도의 규제·윤리·임상 검증 체계에서 검토한다.\n"
    "소감: 1인 개발로 데이터·모델·평가·웹·배포·라이선스까지 연결하면서 높은 평균 점수보다 실패 조건과 재현 근거를 남기는 일이 의료 AI 포트폴리오의 신뢰도를 좌우한다는 점을 확인했다."
)

DATASET = (
    "PhysioNet PTB-XL v1.0.3 (DOI 10.13026/kfzx-aw45), CC BY 4.0. 비식별 환자 18,869명의 12유도 10초 ECG 21,799건 중 "
    "NORM/MI/STTC/CD/HYP 상위군 라벨이 하나 이상인 21,388건을 사용했다. 공식 환자 단위 분할은 학습 17,084건(fold 1-8), "
    "검증 2,146건(fold 9), 시험 2,158건(fold 10)이며 분할 간 환자 중복은 0명이다."
)

PROCESSING = (
    "SCP-ECG 코드를 5개 진단 상위군 다중 라벨로 매핑했다. 5개 라벨이 모두 없는 411건은 메타데이터 감사에는 보존하되 분류 학습·평가에서 제외했다. "
    "100 Hz, 10초, 12유도, (12, 1000), 유한값과 유도 순서를 검증했다. 별도 필터링·재샘플링 없이 학습 fold에서만 계산한 전역 평균·표준편차를 "
    "검증·시험에 고정 적용했다. PTB-XL의 비식별 데이터만 사용했고 추가 개인정보를 수집하지 않았다."
)

AI_DISCLOSURE = (
    "OpenAI Codex를 요구사항 정리, 코드 작성·리팩터링, 테스트·문서화 및 배포 자동화 보조에 활용했다. "
    "데이터 분할, 모델·평가 설계, 수치 해석, 라이선스 확인과 최종 검증 책임은 참가자에게 있으며 생성 결과는 테스트와 재현 절차로 검토했다. "
    "상용 생성형 AI 모델 자체를 ECG 추론 모델이나 학습 데이터로 사용하지 않았다."
)

DEPENDENCIES = [
    ("1", "PyTorch", "2.12.1", "BSD-3-Clause", "https://github.com/pytorch/pytorch", "residual 1D CNN 학습·추론"),
    ("2", "NumPy", "2.5.1", "BSD-3-Clause", "https://github.com/numpy/numpy", "ECG 배열·정규화·수치 연산"),
    ("3", "pandas", "3.0.5", "BSD-3-Clause", "https://github.com/pandas-dev/pandas", "PTB-XL 메타데이터 처리"),
    ("4", "scikit-learn", "1.9.0", "BSD-3-Clause", "https://github.com/scikit-learn/scikit-learn", "AUROC·AUPRC·보정 지표"),
    ("5", "WFDB", "4.3.1", "MIT", "https://github.com/MIT-LCP/wfdb-python", "WFDB 헤더·파형 로딩"),
    ("6", "Matplotlib", "3.11.1", "Matplotlib License", "https://github.com/matplotlib/matplotlib", "ECG·평가 결과 시각화"),
    ("7", "Streamlit", "1.60.0", "Apache-2.0", "https://github.com/streamlit/streamlit", "연구용 웹 데모 UI"),
    ("8", "pytest", "9.1.1", "MIT", "https://github.com/pytest-dev/pytest", "자동화 테스트(개발 전용)"),
]


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor = BLACK) -> None:
    run.font.name = "Malgun Gothic"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Malgun Gothic")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_cell_text(cell, text: str, *, size: float = 10, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_cell_paragraph(cell, text: str, *, label: str | None = None, size: float = 10, after: float = 3) -> None:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if label:
        lead = p.add_run(label)
        set_run(lead, size=size, bold=True)
    body = p.add_run(text)
    set_run(body, size=size)


def clear_cell(cell) -> None:
    for p in list(cell.paragraphs)[1:]:
        p._element.getparent().remove(p._element)
    cell.paragraphs[0].clear()


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def make_architecture_image() -> None:
    width, height = 1800, 460
    image = PILImage.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype(str(MALGUN), 28)
    bold = ImageFont.truetype(str(MALGUN_BOLD), 30)
    small = ImageFont.truetype(str(MALGUN), 23)
    boxes = [
        (40, "PTB-XL", "WFDB 12유도\n100 Hz"),
        (390, "분할·정규화", "환자 단위 fold\n학습 통계 고정"),
        (740, "Residual 1D CNN", "5개 진단군\n다중 라벨 확률"),
        (1090, "신뢰성 평가", "보정·SQI\nuncertainty"),
        (1440, "공개 출력", "CLI·Streamlit\nJSON·모델 카드"),
    ]
    top, bw, bh = 125, 300, 205
    for i, (x, title, subtitle) in enumerate(boxes):
        draw.rounded_rectangle((x, top, x + bw, top + bh), radius=24, fill="#EAF3FB", outline="#4E86B6", width=4)
        title_box = draw.textbbox((0, 0), title, font=bold)
        draw.text((x + (bw - (title_box[2] - title_box[0])) / 2, top + 35), title, font=bold, fill="#173B5E")
        lines = subtitle.split("\n")
        for li, line in enumerate(lines):
            line_box = draw.textbbox((0, 0), line, font=small)
            draw.text((x + (bw - (line_box[2] - line_box[0])) / 2, top + 105 + li * 37), line, font=small, fill="#111111")
        if i < len(boxes) - 1:
            ax1, ax2, ay = x + bw + 10, boxes[i + 1][0] - 15, top + bh // 2
            draw.line((ax1, ay, ax2, ay), fill="#4E86B6", width=6)
            draw.polygon([(ax2, ay), (ax2 - 22, ay - 14), (ax2 - 22, ay + 14)], fill="#4E86B6")
    caption = "데이터 누수 방지 -> 잠금 평가 -> 책임 있는 공개 출력"
    cap_box = draw.textbbox((0, 0), caption, font=font)
    draw.text(((width - (cap_box[2] - cap_box[0])) / 2, 385), caption, font=font, fill="#333333")
    image.save(ARCH_IMAGE)


def build_docx() -> None:
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    if sha256(REFERENCE).upper() != "937679BAC40CBFACED3457530C232C9D190A74F6B5D67C58B4BC33014A579195":
        raise RuntimeError("official template hash changed")
    shutil.copy2(REFERENCE, DOCX_OUT)
    doc = Document(DOCX_OUT)
    source_tables = list(doc.tables)
    guide, title, meta, project, sbom_title, sbom, ai_title, ai_guide, ai = source_tables
    remove_element(guide._element)
    remove_element(ai_guide._element)
    for p in list(doc.paragraphs):
        if p.text.strip().startswith("※ 필요 시, 행을 추가하여"):
            remove_element(p._element)

    set_cell_text(meta.cell(1, 1), TEAM_NAME, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(meta.cell(1, 3), "1명", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(meta.cell(2, 1), "학생", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(meta.cell(2, 3), "자유과제 (사회문제해결: 생활)", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(project.cell(1, 1), PROJECT_NAME, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(project.cell(2, 1), REPOSITORY, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(project.cell(3, 1), "제출 전 YouTube 시연영상 URL 입력 필요", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(project.cell(4, 1), INTRO, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    set_cell_text(project.cell(6, 1), BACKGROUND, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    set_cell_text(project.cell(7, 1), ENVIRONMENT, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    clear_cell(project.cell(8, 1))
    add_cell_paragraph(project.cell(8, 1), ARCHITECTURE, label="구성 및 데이터 흐름. ")
    p = project.cell(8, 1).add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    architecture_shape = p.add_run().add_picture(str(ARCH_IMAGE), width=Inches(4.30))
    architecture_shape._inline.docPr.set("descr", "ECG Guard 데이터 처리와 신뢰성 평가 아키텍처")

    clear_cell(project.cell(9, 1))
    for chunk in FEATURES.split("\n"):
        label, sep, body = chunk.partition(": ")
        add_cell_paragraph(project.cell(9, 1), body if sep else chunk, label=f"{label}. " if sep else None, size=9.6, after=2)
    eval_image = ROOT / "outputs" / "baseline-evaluation" / "evaluation_curves.png"
    p = project.cell(9, 1).add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    evaluation_shape = p.add_run().add_picture(str(eval_image), width=Inches(4.30))
    evaluation_shape._inline.docPr.set("descr", "동결된 baseline-v1의 ROC, precision-recall, calibration 및 하위집단 결과")
    cap = project.cell(9, 1).add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(2)
    set_run(cap.add_run("그림 1. 동결된 baseline-v1 시험 평가 곡선"), size=8.5)

    set_cell_text(project.cell(10, 1), IMPACT, size=9.7, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    clear_cell(project.cell(11, 1))
    for chunk in OTHER.split("\n"):
        label, sep, body = chunk.partition(": ")
        add_cell_paragraph(project.cell(11, 1), body if sep else chunk, label=f"{label}. " if sep else None, size=9.5, after=2)

    while len(sbom.rows) < len(DEPENDENCIES) + 1:
        sbom.add_row()
    for ri, dep in enumerate(DEPENDENCIES, start=1):
        for ci, value in enumerate(dep):
            set_cell_text(sbom.cell(ri, ci), value, size=8.3 if ci in (4, 5) else 8.8, align=WD_ALIGN_PARAGRAPH.CENTER if ci < 4 else WD_ALIGN_PARAGRAPH.LEFT)
            shade_cell(sbom.cell(ri, ci), "FFFFFF")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.1
    set_run(note.add_run("최종 컨테이너 SBOM: "), size=9, bold=True)
    set_run(note.add_run("CycloneDX 1.7, 총 3,079 components(89 PyPI, 88 Debian 포함). "), size=9)
    set_run(note.add_run(f"sbom/container-runtime.cdx.json · SHA-256 {CONTAINER_SBOM_SHA256}"), size=8.5)

    set_cell_text(
        ai.cell(1, 0),
        "□ 유형 1: 외부 모델 그대로 활용\n□ 유형 2: 외부 모델 파인튜닝\n▣ 유형 3: 자체 개발 모델 (residual 1D CNN 전체 가중치를 처음부터 학습)",
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    set_cell_text(ai.cell(3, 1), "해당 없음 (외부 기반 모델을 사용하지 않은 유형 3)", size=9.3)
    set_cell_text(ai.cell(3, 4), "해당 없음", size=9.3)
    set_cell_text(ai.cell(5, 1), DATASET, size=8.9, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    set_cell_text(ai.cell(6, 1), PROCESSING, size=8.9, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    set_cell_text(ai.cell(7, 1), RELEASE_URL, size=8.8)
    set_cell_text(
        ai.cell(8, 1),
        f"파일명: best_model.pt / PyTorch state_dict / 23,579,661 bytes / 인증 없이 공개 다운로드 / SHA-256: {WEIGHT_SHA256}",
        size=8.6,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    set_cell_text(ai.cell(10, 1), "Apache License 2.0 (Apache-2.0)", size=9.3)
    set_cell_text(ai.cell(10, 4), REPOSITORY, size=8.8)
    set_cell_text(ai.cell(11, 1), AI_DISCLOSURE, size=8.9, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    for table in doc.tables:
        mark_header_row(table.rows[0])
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.color.rgb == GRAY:
                            run.font.color.rgb = BLACK

    props = doc.core_properties
    props.title = PROJECT_NAME
    props.subject = "2026 오픈소스 개발자대회 결과보고서"
    props.author = "유라"
    props.keywords = "ECG, medical AI, uncertainty, signal quality, open source"
    props.comments = "Generated from the official contest DOCX template."
    doc.save(DOCX_OUT)
    restore_template_parts(
        DOCX_OUT,
        REFERENCE,
        ["_rels/.rels", "word/numbering.xml", "word/settings.xml", "word/styles.xml"],
    )


def restore_template_parts(final_path: Path, reference_path: Path, names: list[str]) -> None:
    temporary = final_path.with_suffix(".repacked.docx")
    with ZipFile(final_path) as final_zip, ZipFile(reference_path) as reference_zip, ZipFile(
        temporary, "w", ZIP_DEFLATED
    ) as output_zip:
        for info in final_zip.infolist():
            data = reference_zip.read(info.filename) if info.filename in names else final_zip.read(info.filename)
            output_zip.writestr(info, data)
    temporary.replace(final_path)


def esc(text: str) -> str:
    return html.escape(text).replace("\n", "<br/>")


def build_pdf() -> None:
    pdfmetrics.registerFont(TTFont("Malgun", str(MALGUN)))
    pdfmetrics.registerFont(TTFont("MalgunBold", str(MALGUN_BOLD)))
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "KoreanBody",
        parent=styles["BodyText"],
        fontName="Malgun",
        fontSize=9.5,
        leading=13.5,
        alignment=TA_JUSTIFY,
        textColor=colors.black,
        spaceAfter=2,
        wordWrap="CJK",
    )
    body_small = ParagraphStyle("KoreanSmall", parent=body, fontSize=8, leading=10.5)
    label = ParagraphStyle("KoreanLabel", parent=body, fontName="MalgunBold", fontSize=9.5, leading=12, alignment=TA_CENTER)
    section = ParagraphStyle("KoreanSection", parent=label, fontSize=11, leading=14)
    title_style = ParagraphStyle("KoreanTitle", parent=section, fontName="MalgunBold", fontSize=16, leading=20, spaceAfter=8)
    caption = ParagraphStyle("KoreanCaption", parent=body_small, alignment=TA_CENTER, spaceBefore=2, spaceAfter=4)
    url_style = ParagraphStyle("KoreanURL", parent=body_small, alignment=TA_LEFT, fontSize=7.7, leading=9.5)

    def P(text: str, style=body):
        return Paragraph(esc(text), style)

    def on_page(canvas, doc):
        canvas.saveState()
        canvas.setFont("Malgun", 8)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawRightString(doc.pagesize[0] - 15 * mm, 10 * mm, f"{doc.page}")
        canvas.restoreState()

    doc = BaseDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=30 * mm,
        rightMargin=30 * mm,
        topMargin=24 * mm,
        bottomMargin=20 * mm,
        title=PROJECT_NAME,
        author="유라",
    )
    portrait_frame = Frame(30 * mm, 20 * mm, A4[0] - 60 * mm, A4[1] - 44 * mm, id="portrait")
    land_size = landscape(A4)
    landscape_frame = Frame(20 * mm, 18 * mm, land_size[0] - 40 * mm, land_size[1] - 36 * mm, id="landscape")
    doc.addPageTemplates(
        [
            PageTemplate(id="portrait", pagesize=A4, frames=[portrait_frame], onPage=on_page),
            PageTemplate(id="landscape", pagesize=land_size, frames=[landscape_frame], onPage=on_page),
        ]
    )

    grid = colors.HexColor("#4B4B4B")
    blue = colors.HexColor("#A5C9EB")
    pale = colors.HexColor("#DAE9F7")
    story = [P("2026년 오픈소스 개발자대회 결과보고서", title_style)]
    meta = Table(
        [
            [P("항목", label), P("내용", label), P("항목", label), P("내용", label)],
            [P("팀명", label), P(TEAM_NAME, label), P("팀 인원", label), P("1명", label)],
            [P("참가부문", label), P("학생", label), P("과제유형", label), P("자유과제 (사회문제해결: 생활)", label)],
        ],
        colWidths=[28 * mm, 47 * mm, 28 * mm, 47 * mm],
        repeatRows=1,
    )
    meta.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.6, grid),
                ("BACKGROUND", (0, 0), (-1, 0), blue),
                ("BACKGROUND", (0, 1), (0, -1), pale),
                ("BACKGROUND", (2, 1), (2, -1), pale),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story += [meta, Spacer(1, 5 * mm)]

    arch_flow = [P(ARCHITECTURE, body), Spacer(1, 2 * mm), Image(str(ARCH_IMAGE), width=110 * mm, height=28 * mm)]
    eval_image = Image(str(ROOT / "outputs" / "baseline-evaluation" / "evaluation_curves.png"), width=110 * mm, height=70.5 * mm)
    feature_flow = [P(text, body) for text in FEATURES.split("\n")]
    feature_flow += [Spacer(1, 2 * mm), eval_image, P("그림 1. 동결된 baseline-v1 시험 평가 곡선", caption)]
    project_data = [
        [P("프로젝트 개요", section), ""],
        [P("프로젝트명", label), P(PROJECT_NAME, body)],
        [P("프로젝트 등록 URL", label), P(REPOSITORY, url_style)],
        [P("시연영상", label), P("제출 전 YouTube 시연영상 URL 입력 필요", body)],
        [P("프로젝트 소개", label), P(INTRO, body)],
        [P("프로젝트 세부 내용", section), ""],
        [P("개발배경 및 목적", label), P(BACKGROUND, body)],
        [P("개발환경", label), P(ENVIRONMENT, body)],
        [P("시스템 구성 및 아키텍처", label), arch_flow],
        [P("프로젝트 주요기능", label), feature_flow],
        [P("기대효과 및 활용분야", label), P(IMPACT, body)],
        [P("기타", label), [P(text, body) for text in OTHER.split("\n")]],
    ]
    project_table = Table(project_data, colWidths=[35 * mm, 115 * mm], repeatRows=0, splitByRow=1)
    project_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.55, grid),
                ("SPAN", (0, 0), (-1, 0)),
                ("SPAN", (0, 5), (-1, 5)),
                ("BACKGROUND", (0, 0), (-1, 0), blue),
                ("BACKGROUND", (0, 5), (-1, 5), blue),
                ("BACKGROUND", (0, 1), (0, 4), pale),
                ("BACKGROUND", (0, 6), (0, -1), pale),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story += [project_table, NextPageTemplate("landscape"), PageBreak()]

    story += [P("붙임1  SBOM(소프트웨어 자재명세서)", title_style)]
    dep_data = [[P(x, label) for x in ["번호", "라이브러리명", "버전", "라이선스", "공식 저장소 URL", "사용 목적 및 주요 기능"]]]
    for dep in DEPENDENCIES:
        dep_data.append([P(dep[0], body_small), P(dep[1], body_small), P(dep[2], body_small), P(dep[3], body_small), P(dep[4], url_style), P(dep[5], body_small)])
    dep_table = Table(dep_data, colWidths=[12 * mm, 31 * mm, 19 * mm, 26 * mm, 75 * mm, 94 * mm], repeatRows=1)
    dep_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, grid),
                ("BACKGROUND", (0, 0), (-1, 0), pale),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 1), (3, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story += [dep_table, Spacer(1, 4 * mm), P(f"최종 컨테이너 SBOM: CycloneDX 1.7, 총 3,079 components(89 PyPI, 88 Debian 포함). sbom/container-runtime.cdx.json · SHA-256 {CONTAINER_SBOM_SHA256}", body_small), PageBreak()]

    story += [P("붙임2  AI 모델 활용 및 라이선스 기술 명세서", title_style)]
    ai_data = [
        [P("1. AI 모델 활용 유형", section), "", "", "", ""],
        [P("□ 유형 1: 외부 모델 그대로 활용\n□ 유형 2: 외부 모델 파인튜닝\n▣ 유형 3: 자체 개발 모델 (residual 1D CNN 전체 가중치를 처음부터 학습)", body), "", "", "", ""],
        [P("2. 기반(베이스) 모델 정보", section), "", "", "", ""],
        [P("기반 모델명 및 개발사", label), P("해당 없음 (유형 3)", body), "", P("기반 모델 라이선스", label), P("해당 없음", body)],
        [P("3. 데이터셋 정보 및 가중치 배포 명세", section), "", "", "", ""],
        [P("학습 데이터셋 정보", label), P(DATASET, body), "", "", ""],
        [P("데이터 정제/가공 방법", label), P(PROCESSING, body), "", "", ""],
        [P("가중치 공개 URL", label), P(RELEASE_URL, url_style), "", "", ""],
        [P("가중치 파일 정보", label), P(f"best_model.pt / PyTorch state_dict / 23,579,661 bytes / 공개 다운로드 / SHA-256 {WEIGHT_SHA256}", body_small), "", "", ""],
        [P("4. 소스코드 라이선스 및 개발 환경", section), "", "", "", ""],
        [P("코드 라이선스", label), P("Apache License 2.0", body), "", P("소스코드 저장소", label), P(REPOSITORY, url_style)],
        [P("상용 AI 보조도구 활용", label), P(AI_DISCLOSURE, body), "", "", ""],
    ]
    ai_table = Table(ai_data, colWidths=[43 * mm, 87 * mm, 3 * mm, 48 * mm, 76 * mm], splitByRow=1)
    ai_style = [
        ("GRID", (0, 0), (-1, -1), 0.55, grid),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    for row in (0, 1, 2, 4, 5, 6, 7, 8, 9, 11):
        if row in (0, 1, 2, 4, 9):
            ai_style.append(("SPAN", (0, row), (-1, row)))
    for row in (5, 6, 7, 8, 11):
        ai_style.append(("SPAN", (1, row), (-1, row)))
    for row in (0, 2, 4, 9):
        ai_style.append(("BACKGROUND", (0, row), (-1, row), blue))
    for cell in ((0, 3), (3, 3), (0, 5), (0, 6), (0, 7), (0, 8), (0, 10), (3, 10), (0, 11)):
        ai_style.append(("BACKGROUND", cell, cell, pale))
    ai_table.setStyle(TableStyle(ai_style))
    story += [ai_table]
    doc.build(story)


def main() -> None:
    make_architecture_image()
    build_docx()
    build_pdf()
    print(f"DOCX={DOCX_OUT}")
    print(f"PDF={PDF_OUT}")
    print(f"DOCX_SHA256={sha256(DOCX_OUT)}")
    print(f"PDF_SHA256={sha256(PDF_OUT)}")


if __name__ == "__main__":
    main()
